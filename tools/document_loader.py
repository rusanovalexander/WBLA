"""
Document Loading Tools for Credit Pack PoC v2.

Supports:
- TXT files (primary - for anonymized documents)
- PDF (with optional OCR via Document AI)
- DOCX (Word documents)
- XLSX (Excel spreadsheets)
"""

import os
from pathlib import Path
from typing import Dict, Any, List
import glob


from config.settings import (
    PROJECT_ID, LOCATION, DOCAI_PROCESSOR_ID,
    DATA_FOLDER, get_credentials
)


# =============================================================================
# Core Loaders
# =============================================================================

def load_text_file(file_path: str) -> str:
    """Load plain text file (TXT, MD)."""
    try:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    except Exception as e:
        return f"[ERROR] Failed to load text file: {e}"


def load_pdf_simple(file_path: str) -> str:
    """Extract text from PDF using PyMuPDF (no OCR)."""
    try:
        import fitz  # PyMuPDF
        
        doc = fitz.open(file_path)
        text_parts = []
        for page_num, page in enumerate(doc):
            text = page.get_text()
            if text.strip():
                text_parts.append(f"--- Page {page_num + 1} ---\n{text}")
        doc.close()
        
        return '\n\n'.join(text_parts)
    except ImportError:
        return "[ERROR] PyMuPDF not installed. Run: pip install PyMuPDF"
    except Exception as e:
        return f"[PDF_ERROR] {e}"


def load_pdf_with_docai(file_path: str, max_pages: int = 15) -> str:
    """Extract text from PDF using Document AI OCR."""
    try:
        from google.cloud import documentai
        from google.api_core.client_options import ClientOptions
        import fitz
        import tempfile
        
        if not DOCAI_PROCESSOR_ID:
            return "[ERROR] DOCAI_PROCESSOR_ID not configured"
        
        # Check page count
        doc = fitz.open(file_path)
        total_pages = len(doc)
        doc.close()
        
        if total_pages <= max_pages:
            return _docai_single_request(file_path)
        else:
            return _docai_chunked_request(file_path, total_pages, max_pages)
            
    except ImportError as e:
        return f"[ERROR] Missing dependency: {e}"
    except Exception as e:
        return f"[DOCAI_ERROR] {e}"


def _docai_single_request(file_path: str) -> str:
    """Process single document with DocAI."""
    from google.cloud import documentai
    from google.api_core.client_options import ClientOptions
    
    opts = ClientOptions(api_endpoint=f"{LOCATION}-documentai.googleapis.com")
    client = documentai.DocumentProcessorServiceClient(
        client_options=opts,
        credentials=get_credentials()
    )
    
    processor_name = client.processor_path(PROJECT_ID, LOCATION, DOCAI_PROCESSOR_ID)
    
    with open(file_path, "rb") as f:
        content = f.read()
    
    ext = Path(file_path).suffix.lower()
    mime_types = {
        ".pdf": "application/pdf",
        ".png": "image/png",
        ".jpg": "image/jpeg",
        ".jpeg": "image/jpeg",
    }
    mime_type = mime_types.get(ext, "application/pdf")
    
    raw_document = documentai.RawDocument(content=content, mime_type=mime_type)
    request = documentai.ProcessRequest(name=processor_name, raw_document=raw_document)
    
    result = client.process_document(request=request)
    return result.document.text


def _docai_chunked_request(file_path: str, total_pages: int, max_pages: int) -> str:
    """Process large PDF in chunks."""
    import fitz
    import tempfile
    
    all_text = []
    doc = fitz.open(file_path)
    
    for chunk_start in range(0, total_pages, max_pages):
        chunk_end = min(chunk_start + max_pages, total_pages)
        
        # Create temp PDF with chunk
        temp_doc = fitz.open()
        for page_num in range(chunk_start, chunk_end):
            temp_doc.insert_pdf(doc, from_page=page_num, to_page=page_num)
        
        with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as tmp:
            temp_doc.save(tmp.name)
            tmp_path = tmp.name
        temp_doc.close()
        
        try:
            chunk_text = _docai_single_request(tmp_path)
            all_text.append(f"--- Pages {chunk_start+1}-{chunk_end} ---\n{chunk_text}")
        finally:
            os.unlink(tmp_path)
    
    doc.close()
    return '\n\n'.join(all_text)


def load_docx(file_path: str) -> str:
    """Extract text from Word document."""
    try:
        from docx import Document
        
        doc = Document(file_path)
        text_parts = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        # Extract tables
        for table in doc.tables:
            table_rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                table_rows.append(" | ".join(cells))
            if table_rows:
                text_parts.append("\n[TABLE]\n" + "\n".join(table_rows) + "\n[/TABLE]")
        
        return '\n\n'.join(text_parts)
    except ImportError:
        return "[ERROR] python-docx not installed"
    except Exception as e:
        return f"[DOCX_ERROR] {e}"


def load_excel(file_path: str) -> str:
    """Extract text from Excel file as markdown tables."""
    try:
        import pandas as pd
        
        xls = pd.ExcelFile(file_path)
        text_parts = [f"--- FILE: {os.path.basename(file_path)} ---"]
        
        for sheet_name in xls.sheet_names:
            df = pd.read_excel(xls, sheet_name=sheet_name)
            text_parts.append(f"\n### SHEET: {sheet_name}\n")
            text_parts.append(df.to_markdown(index=False))
        
        return '\n'.join(text_parts)
    except ImportError:
        return "[ERROR] pandas not installed"
    except Exception as e:
        return f"[EXCEL_ERROR] {e}"


# =============================================================================
# Universal Loader
# =============================================================================

def universal_loader(file_path: str, force_ocr: bool = False) -> str:
    """
    Smart document loader - routes to appropriate handler.
    
    Args:
        file_path: Path to document
        force_ocr: Force OCR for PDFs (for scanned documents)
    
    Returns:
        Extracted text
    """
    if not os.path.exists(file_path):
        return f"[ERROR] File not found: {file_path}"
    
    ext = Path(file_path).suffix.lower()
    
    # Text files (primary for anonymized docs)
    if ext in [".txt", ".md"]:
        return load_text_file(file_path)
    
    # PDF
    elif ext == ".pdf":
        if force_ocr:
            return load_pdf_with_docai(file_path)
        else:
            # Try simple extraction first
            text = load_pdf_simple(file_path)
            if len(text.strip()) < 100:  # Likely scanned
                return load_pdf_with_docai(file_path)
            return text
    
    # Word
    elif ext == ".docx":
        return load_docx(file_path)
    
    # Excel
    elif ext in [".xlsx", ".xls"]:
        return load_excel(file_path)
    
    # Images (OCR)
    elif ext in [".png", ".jpg", ".jpeg"]:
        return load_pdf_with_docai(file_path)
    
    else:
        return f"[ERROR] Unsupported format: {ext}"


# =============================================================================
# Folder Scanner
# =============================================================================

def scan_data_folder() -> Dict[str, List[str]]:
    """
    Scan data folder and categorize files.
    
    Returns:
        Dict with categorized file paths
    """
    categories = {
        "teasers": [],
        "examples": [],
        "procedure": [],
        "guidelines": [],
        "other": [],
    }
    
    # Check subfolders
    for category in ["teasers", "examples", "procedure", "guidelines"]:
        folder = DATA_FOLDER / category
        if folder.exists():
            for f in folder.iterdir():
                if f.is_file() and not f.name.startswith('.'):
                    categories[category].append(str(f))
    
    # Also check root data folder
    for f in DATA_FOLDER.iterdir():
        if f.is_file() and not f.name.startswith('.'):
            fname = f.name.lower()
            if "teaser" in fname:
                categories["teasers"].append(str(f))
            elif "example" in fname or "credit_pack" in fname or "creditpack" in fname:
                categories["examples"].append(str(f))
            elif "procedure" in fname:
                categories["procedure"].append(str(f))
            elif "guideline" in fname or "guidance" in fname:
                categories["guidelines"].append(str(f))
            else:
                categories["other"].append(str(f))
    
    return categories


# =============================================================================
# Tool Wrappers (for agents)
# =============================================================================

def tool_load_document(file_path: str, force_ocr: bool = False) -> Dict[str, Any]:
    """
    Load a document and extract text.
    
    Args:
        file_path: Path to document
        force_ocr: Whether to force OCR processing
        
    Returns:
        Dict with status, file info, and extracted text
    """
    if not os.path.exists(file_path):
        return {
            "status": "ERROR",
            "error": f"File not found: {file_path}",
            "text": ""
        }
    
    text = universal_loader(file_path, force_ocr=force_ocr)
    
    if text.startswith("[ERROR]"):
        return {
            "status": "ERROR",
            "error": text,
            "text": ""
        }
    
    # Truncate if very long
    max_length = 150000
    truncated = False
    if len(text) > max_length:
        text = text[:max_length] + "\n\n[TRUNCATED]"
        truncated = True
    
    return {
        "status": "OK",
        "file_path": file_path,
        "file_name": os.path.basename(file_path),
        "file_type": Path(file_path).suffix.lower(),
        "text_length": len(text),
        "truncated": truncated,
        "text": text
    }


def tool_scan_data_folder() -> Dict[str, Any]:
    """
    Scan data folder and return categorized file list.
    
    Returns:
        Dict with folder contents summary
    """
    categories = scan_data_folder()
    
    summary = {cat: len(files) for cat, files in categories.items()}
    
    return {
        "status": "OK",
        "folder": str(DATA_FOLDER),
        "summary": summary,
        "files": categories
    }


def tool_load_teaser() -> Dict[str, Any]:
    """Load the first teaser document found."""
    categories = scan_data_folder()
    
    if not categories["teasers"]:
        return {"status": "ERROR", "error": "No teaser found in data/teasers/"}
    
    return tool_load_document(categories["teasers"][0], force_ocr=True)


def tool_load_example() -> Dict[str, Any]:
    """Load the first example credit pack found."""
    categories = scan_data_folder()
    
    if not categories["examples"]:
        return {"status": "ERROR", "error": "No example found in data/examples/"}
    
    return tool_load_document(categories["examples"][0])


# =============================================================================
# Test
# =============================================================================

if __name__ == "__main__":
    print("Testing document loader...")
    
    result = tool_scan_data_folder()
    print(f"\nFolder scan:")
    for cat, count in result["summary"].items():
        print(f"  {cat}: {count} files")
        for f in result["files"][cat][:2]:
            print(f"    - {os.path.basename(f)}")
