"""DOCX export. Self-contained, uses python-docx."""

import logging
import re
from pathlib import Path

from .config import OUTPUTS_DIR, PRODUCT_NAME

logger = logging.getLogger(__name__)


def generate_docx(content: str, filename: str, metadata: dict | None = None) -> str:
    """Generate DOCX from markdown-style content. Returns path to saved file or empty string."""
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        Path(OUTPUTS_DIR).mkdir(parents=True, exist_ok=True)
        out_path = Path(OUTPUTS_DIR) / filename
        doc = Document()
        section = doc.sections[0]
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

        for para in doc.paragraphs:
            para.paragraph_format.space_after = Pt(6)

        # Simple markdown-to-docx: split by # headings
        blocks = re.split(r"\n(?=# )", content)
        for block in blocks:
            block = block.strip()
            if not block:
                continue
            if block.startswith("# "):
                doc.add_heading(block[2:].split("\n")[0], level=1)
                rest = "\n".join(block.split("\n")[1:]).strip()
                if rest:
                    doc.add_paragraph(rest)
            else:
                doc.add_paragraph(block)

        doc.save(str(out_path))
        return str(out_path)
    except Exception as e:
        logger.exception("generate_docx failed: %s", e)
        return ""
