"""
Export module for Credit Pack PoC v3.2.

Generates professional DOCX output and audit trail files.
"""

from __future__ import annotations

import logging
import re
from datetime import datetime
from pathlib import Path
from typing import Any

from config.settings import OUTPUTS_FOLDER, VERSION
from core.tracing import TraceStore, get_tracer

logger = logging.getLogger(__name__)


# =============================================================================
# DOCX Generation — Professional Banking Format
# =============================================================================

def generate_docx(
    content: str,
    filename: str,
    metadata: dict[str, str] | None = None,
) -> str:
    """
    Generate professional DOCX from markdown content.

    Args:
        content: Markdown content (section bodies joined by ---)
        filename: Output filename
        metadata: Optional metadata (deal_name, borrower, process_path, etc.)

    Returns:
        Path to saved file, or "" on error
    """
    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        doc = Document()

        # ---- Page Setup ----
        section = doc.sections[0]
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

        # ---- Styles ----
        normal = doc.styles["Normal"]
        normal.font.name = "Calibri"
        normal.font.size = Pt(10.5)
        normal.paragraph_format.space_after = Pt(6)
        normal.paragraph_format.line_spacing = 1.15

        for level, (size, bold) in enumerate([(18, True), (14, True), (12, True)], 1):
            hs = doc.styles[f"Heading {level}"]
            hs.font.name = "Calibri"
            hs.font.size = Pt(size)
            hs.font.bold = bold
            hs.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
            hs.paragraph_format.space_before = Pt(18 if level == 1 else 12)
            hs.paragraph_format.space_after = Pt(8)

        # ---- Cover Page ----
        doc.add_paragraph()
        doc.add_paragraph()
        doc.add_paragraph()

        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run("CREDIT PACK")
        title_run.font.size = Pt(28)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
        title_run.font.name = "Calibri"

        # Deal name: prefer analysis-derived name over raw filename
        deal_name = ""
        if metadata:
            # Try structured deal info first
            deal_name = (
                metadata.get("borrower_name")
                or metadata.get("deal_name", "")
            )
            # Strip raw filename extension if that's all we have
            if deal_name and deal_name.endswith((".txt", ".pdf", ".docx")):
                deal_name = Path(deal_name).stem.replace("_", " ").replace("-", " ").title()

        if deal_name:
            subtitle = doc.add_paragraph()
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sub_run = subtitle.add_run(deal_name)
            sub_run.font.size = Pt(16)
            sub_run.font.color.rgb = RGBColor(0x4A, 0x4A, 0x4A)
            sub_run.font.name = "Calibri"

        doc.add_paragraph()

        # Cover details — assessment, origination, date only (no system version)
        cover_items: list[tuple[str, str]] = [
            ("Date", datetime.now().strftime("%d %B %Y")),
        ]
        if metadata:
            if metadata.get("process_path"):
                cover_items.append(("Assessment Approach", metadata["process_path"]))
            if metadata.get("origination_method"):
                cover_items.append(("Origination Method", metadata["origination_method"]))
            if metadata.get("borrower_name"):
                cover_items.insert(0, ("Borrower", metadata["borrower_name"]))

        for label, value in cover_items:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            label_run = p.add_run(f"{label}: ")
            label_run.font.size = Pt(10)
            label_run.font.bold = True
            label_run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
            val_run = p.add_run(value)
            val_run.font.size = Pt(10)
            val_run.font.color.rgb = RGBColor(0x4A, 0x4A, 0x4A)

        doc.add_paragraph()
        _add_classification_banner(doc, "CONFIDENTIAL — FOR INTERNAL USE ONLY")

        # Page break after cover
        doc.add_page_break()

        # ---- Table of Contents ----
        toc_heading = doc.add_heading("Contents", level=1)
        toc_heading.paragraph_format.space_before = Pt(0)
        _add_table_of_contents(doc)
        doc.add_page_break()

        # ---- Content (with section numbering) ----
        _render_markdown_to_docx(doc, content, number_h1=True)

        # ---- Footer ----
        footer = doc.sections[0].footer
        footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer_para.add_run(
            f"Credit Pack — {deal_name} — "
            f"Generated {datetime.now().strftime('%d %B %Y')} — CONFIDENTIAL"
        )
        footer_run.font.size = Pt(8)
        footer_run.font.color.rgb = RGBColor(0xA0, 0xA0, 0xA0)

        # ---- Save ----
        output_dir = Path(OUTPUTS_FOLDER)
        output_dir.mkdir(parents=True, exist_ok=True)
        output_path = output_dir / filename
        doc.save(str(output_path))

        logger.info("DOCX saved: %s", output_path)
        return str(output_path)

    except ImportError:
        logger.error("python-docx not installed")
        return ""
    except Exception as e:
        logger.error("DOCX generation failed: %s", e, exc_info=True)
        return ""


def _add_table_of_contents(doc) -> None:
    """
    Insert a Word TOC field that auto-updates when the document is opened.
    Word renders this as a proper clickable Table of Contents.
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fldChar_begin = OxmlElement("w:fldChar")
    fldChar_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fldChar_begin)

    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    run._r.append(instrText)

    fldChar_separate = OxmlElement("w:fldChar")
    fldChar_separate.set(qn("w:fldCharType"), "separate")
    run._r.append(fldChar_separate)

    fldChar_end = OxmlElement("w:fldChar")
    fldChar_end.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar_end)

    # Instruction to user
    note = doc.add_paragraph()
    note_run = note.add_run("(Right-click → Update Field to refresh the Table of Contents)")
    note_run.font.italic = True
    note_run.font.size = __import__("docx.shared", fromlist=["Pt"]).Pt(9)
    from docx.shared import RGBColor
    note_run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)


def _add_classification_banner(doc, text: str):
    """Add a classification banner paragraph."""
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)


def _normalise_tables(content: str) -> str:
    """Convert tab-separated table rows to markdown pipe format (safety net)."""
    normalised = []
    for line in content.split("\n"):
        if "\t" in line and not line.strip().startswith("|"):
            cells = [c.strip() for c in line.split("\t")]
            if len(cells) >= 2:
                line = "| " + " | ".join(cells) + " |"
        normalised.append(line)
    return "\n".join(normalised)


def _render_markdown_to_docx(doc, content: str, number_h1: bool = False):
    """Convert markdown content to DOCX elements with proper formatting.

    Args:
        doc: python-docx Document object
        content: Markdown string to render
        number_h1: If True, prefix every H1 heading with a sequential section number
                   (e.g. "1. Executive Summary", "2. Borrower Overview", …)
    """
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    content = _normalise_tables(content)
    lines = content.split("\n")
    current_table_rows: list[list[str]] = []
    in_table = False
    h1_counter = 0  # section numbering counter

    for line in lines:
        stripped = line.strip()

        # Table detection
        if stripped.startswith("|") and "|" in stripped[1:]:
            cells = [c.strip() for c in stripped.split("|")[1:-1]]
            # Skip separator rows
            if cells and not all(c.replace("-", "").replace(":", "").strip() == "" for c in cells):
                current_table_rows.append(cells)
                in_table = True
            continue
        elif in_table and current_table_rows:
            _add_styled_table(doc, current_table_rows)
            current_table_rows = []
            in_table = False

        # Skip empty lines
        if not stripped:
            continue

        # Headings — full hierarchy
        if stripped.startswith("# ") and not stripped.startswith("## "):
            heading_text = stripped[2:].strip()
            if number_h1:
                h1_counter += 1
                heading_text = f"{h1_counter}. {heading_text}"
            doc.add_heading(heading_text, level=1)
        elif stripped.startswith("## ") and not stripped.startswith("### "):
            doc.add_heading(stripped[3:].strip(), level=2)
        elif stripped.startswith("### ") and not stripped.startswith("#### "):
            doc.add_heading(stripped[4:].strip(), level=3)
        elif stripped.startswith("#### "):
            doc.add_heading(stripped[5:].strip(), level=3)
        elif stripped.startswith("---"):
            # Horizontal rule — subtle spacing
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            # Bullet point
            bullet_text = stripped[2:].strip()
            try:
                p = doc.add_paragraph(style="List Bullet")
            except KeyError:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.5)
                # Add bullet character manually
                bullet_text = f"• {bullet_text}"
            _add_formatted_text(p, bullet_text)
        elif len(stripped) > 2 and stripped[0].isdigit() and ". " in stripped[:5]:
            # Numbered list (e.g., "1. Item")
            num_prefix = stripped.split(". ", 1)[0]
            num_text = stripped.split(". ", 1)[1] if ". " in stripped else stripped
            try:
                p = doc.add_paragraph(style="List Number")
                _add_formatted_text(p, num_text)
            except KeyError:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.5)
                _add_formatted_text(p, f"{num_prefix}. {num_text}")
        else:
            # Regular paragraph with inline formatting
            p = doc.add_paragraph()
            _add_formatted_text(p, stripped)

    # Flush remaining table
    if current_table_rows:
        _add_styled_table(doc, current_table_rows)


def _add_formatted_text(paragraph, text: str):
    """Add text to paragraph with bold/italic formatting preserved.

    [INFORMATION REQUIRED: …] markers are rendered in red bold with a yellow
    highlight background so they stand out clearly in the printed document.
    """
    from docx.shared import RGBColor
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    # Handle [INFORMATION REQUIRED: ...] markers
    if "[INFORMATION REQUIRED" in text:
        # Match both [INFORMATION REQUIRED: ...] and bare [INFORMATION REQUIRED]
        parts = re.split(r"(\[INFORMATION REQUIRED[^\]]*\])", text)
        for part in parts:
            if part.startswith("[INFORMATION REQUIRED"):
                run = paragraph.add_run(part)
                run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
                run.font.bold = True
                # --- Yellow highlight via w:rPr / w:highlight ---
                rPr = run._r.get_or_add_rPr()
                highlight = OxmlElement("w:highlight")
                highlight.set(qn("w:val"), "yellow")
                rPr.append(highlight)
            else:
                _add_inline_formatted(paragraph, part)
        return

    _add_inline_formatted(paragraph, text)


def _add_inline_formatted(paragraph, text: str):
    """Handle **bold** and *italic* inline formatting."""
    # Split on bold markers
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.font.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.font.italic = True
        elif part:
            # Clean remaining markdown (preserve whitespace-only parts for word spacing)
            clean = re.sub(r"\*([^*]+)\*", r"\1", part)
            paragraph.add_run(clean)


def _add_styled_table(doc, rows: list[list[str]]):
    """Add a professionally styled table to the document."""
    from docx.shared import Pt, Inches, RGBColor, Cm
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn

    if not rows:
        return

    # Normalize column count
    max_cols = max(len(row) for row in rows)
    normalized = [row + [""] * (max_cols - len(row)) for row in rows]

    table = doc.add_table(rows=len(normalized), cols=max_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, row_data in enumerate(normalized):
        for j, cell_text in enumerate(row_data):
            cell = table.cell(i, j)
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(cell_text.strip())
            run.font.size = Pt(9)
            run.font.name = "Calibri"

            # Header row styling
            if i == 0:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                shading = cell._element.get_or_add_tcPr()
                shading_elm = shading.makeelement(
                    qn("w:shd"),
                    {qn("w:val"): "clear", qn("w:color"): "auto", qn("w:fill"): "1B3A5C"}
                )
                shading.append(shading_elm)
            elif i % 2 == 0:
                # Alternating row color
                shading = cell._element.get_or_add_tcPr()
                shading_elm = shading.makeelement(
                    qn("w:shd"),
                    {qn("w:val"): "clear", qn("w:color"): "auto", qn("w:fill"): "F2F7FB"}
                )
                shading.append(shading_elm)

    doc.add_paragraph()  # Spacing after table


# =============================================================================
# Audit Trail Export
# =============================================================================

def generate_audit_trail(
    session_state: dict[str, Any],
    tracer: TraceStore | None = None,
    filename: str = "",
) -> str:
    """
    Generate comprehensive audit trail file.

    Collects all in-memory audit data: process decisions, agent activity,
    inter-agent communication, human changes, phase history, RAG sources.
    """
    if tracer is None:
        tracer = get_tracer()

    if not filename:
        filename = f"audit_trail_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt"

    lines = [
        "=" * 70,
        f"CREDIT PACK AUDIT TRAIL — v{VERSION}",
        f"Generated: {datetime.now().isoformat()}",
        "=" * 70,
        "",
    ]

    # 1. Process Decision
    decision = session_state.get("process_decision")
    if decision:
        lines.extend([
            "─" * 50,
            "PROCESS DECISION",
            "─" * 50,
            f"Assessment Approach: {decision.get('assessment_approach', 'N/A')}",
            f"Origination Method: {decision.get('origination_method', 'N/A')}",
            f"Locked: {decision.get('locked', False)}",
            f"Timestamp: {decision.get('timestamp', 'N/A')}",
        ])
        evidence = decision.get("evidence", {})
        if evidence:
            lines.append(f"Deal Size: {evidence.get('deal_size', 'Unknown')}")
            lines.append(f"Reasoning: {evidence.get('reasoning', 'N/A')}")
        lines.append("")

    # 2. Agent Activity Trace
    lines.extend([
        "─" * 50,
        "AGENT ACTIVITY TRACE",
        "─" * 50,
        tracer.format_for_export(),
        "",
    ])

    # 3. Agent Summary
    summary = tracer.get_agent_summary()
    if summary:
        lines.extend(["─" * 50, "AGENT COST SUMMARY", "─" * 50])
        for agent, stats in summary.items():
            lines.append(
                f"  {agent:25s} | Calls: {stats['calls']:2d} | "
                f"Tokens: {stats['tokens_in']:,} in / {stats['tokens_out']:,} out | "
                f"Cost: ${stats['cost_usd']:.4f} | Time: {stats['total_ms']}ms"
            )
        tokens_in, tokens_out = tracer.total_tokens
        lines.extend([
            "",
            f"  TOTAL: {tracer.total_calls} calls | "
            f"{tokens_in:,} in / {tokens_out:,} out | ${tracer.total_cost:.4f}",
            "",
        ])

    # 4. Change Log
    change_log = session_state.get("change_log")
    if change_log and hasattr(change_log, "has_changes") and change_log.has_changes():
        lines.extend(["─" * 50, "HUMAN CHANGES", "─" * 50])
        lines.append(change_log.generate_audit_trail())
        lines.append("")

    # 5. Agent Communication
    agent_bus = session_state.get("agent_bus")
    if agent_bus and hasattr(agent_bus, "message_count") and agent_bus.message_count > 0:
        lines.extend([
            "─" * 50,
            "AGENT-TO-AGENT COMMUNICATION",
            "─" * 50,
            agent_bus.get_log_formatted(),
            "",
        ])

    # 6. Phase History
    phase_manager = session_state.get("phase_manager")
    if phase_manager and hasattr(phase_manager, "get_phase_history"):
        history = phase_manager.get_phase_history()
        if history:
            lines.extend(["─" * 50, "PHASE NAVIGATION HISTORY", "─" * 50])
            for h in history:
                nav = h.get("navigation", "forward")
                lines.append(f"  {h.get('timestamp', '')} | {h['from']} → {h['to']} ({nav})")
            lines.append("")

    # Save to outputs folder
    output_dir = Path(OUTPUTS_FOLDER)
    output_dir.mkdir(parents=True, exist_ok=True)
    output_path = output_dir / filename

    with open(output_path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))

    logger.info("Audit trail saved: %s", output_path)
    return str(output_path)
